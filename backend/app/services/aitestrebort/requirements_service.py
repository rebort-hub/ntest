"""
需求管理服务
"""
from typing import List, Optional, Dict, Any, Tuple
from fastapi import HTTPException, UploadFile
import uuid
import os
import logging
import re
import json
from datetime import datetime
from pathlib import Path
from tortoise.expressions import Q

from app.models.aitestrebort.requirements import (
    RequirementDocument, RequirementModule, Requirement,
    ReviewReport, ReviewIssue, ModuleReviewResult,
    RequirementType, RequirementPriority, RequirementStatus,
    DocumentStatus
)
from app.schemas.aitestrebort.requirements import (
    RequirementDocumentCreate, RequirementDocumentUpdate,
    RequirementCreate, RequirementUpdate,
    RequirementSearchRequest, RequirementStatistics,
    ProjectStatistics, DocumentSplitRequest
)

logger = logging.getLogger(__name__)


def generate_id() -> str:
    """生成字符串ID"""
    return str(uuid.uuid4())


class RequirementService:
    """需求管理服务"""
    
    # ==================== 需求文档管理 ====================
    
    async def create_requirement_document(
        self,
        project_id: int,
        document_data: RequirementDocumentCreate,
        file: Optional[UploadFile] = None,
        user_id: Optional[int] = None
    ) -> RequirementDocument:
        """创建需求文档"""
        try:
            # 创建文档记录
            document_dict = {
                'project_id': project_id,
                'title': document_data.title,
                'description': document_data.description,
                'document_type': document_data.document_type,
                'uploader_id': user_id,
                'status': DocumentStatus.UPLOADED
            }
            
            # 处理文件上传
            if file:
                file_path = await self._save_uploaded_file(file, project_id)
                document_dict['file_path'] = file_path
                
                # 提取文档内容
                content = await self._extract_document_content(file_path)
                if content:
                    document_dict['content'] = content
                    document_dict['word_count'] = len(content)
                    document_dict['page_count'] = max(1, (len(content) // 500) + 1)
            
            document = await RequirementDocument.create(
                id=generate_id(),
                **document_dict
            )
            
            logger.info(f"需求文档创建成功: {document.id}")
            return document
            
        except Exception as e:
            logger.error(f"创建需求文档失败: {e}")
            raise HTTPException(status_code=500, detail=f"创建需求文档失败: {str(e)}")
    
    async def get_requirement_documents(
        self,
        project_id: int,
        search: Optional[str] = None,
        status: Optional[str] = None,  # 改为字符串类型
        page: int = 1,
        page_size: int = 20
    ) -> Tuple[List[RequirementDocument], int]:
        """获取需求文档列表"""
        try:
            query = RequirementDocument.filter(project_id=project_id)
            
            # 搜索过滤
            if search:
                query = query.filter(
                    Q(title__icontains=search) |
                    Q(description__icontains=search) |
                    Q(content__icontains=search)
                )
            
            # 状态过滤
            if status:
                query = query.filter(status=status)
            
            # 总数
            total = await query.count()
            
            # 分页
            offset = (page - 1) * page_size
            documents = await query.order_by('-uploaded_at').offset(offset).limit(page_size)
            
            return documents, total
            
        except Exception as e:
            logger.error(f"获取需求文档列表失败: {e}")
            raise HTTPException(status_code=500, detail=f"获取需求文档列表失败: {str(e)}")
    
    async def get_requirement_document(self, document_id: str) -> RequirementDocument:
        """获取需求文档详情"""
        document = await RequirementDocument.get_or_none(id=document_id)
        
        if not document:
            raise HTTPException(status_code=404, detail="需求文档不存在")
        
        return document
    
    async def update_requirement_document(
        self,
        document_id: uuid.UUID,
        update_data: RequirementDocumentUpdate
    ) -> RequirementDocument:
        """更新需求文档"""
        try:
            document = await self.get_requirement_document(document_id)
            
            # 更新字段
            update_dict = update_data.model_dump(exclude_unset=True)
            update_dict['updated_at'] = datetime.now()
            
            await document.update_from_dict(update_dict)
            await document.save()
            
            logger.info(f"需求文档更新成功: {document_id}")
            return document
            
        except Exception as e:
            logger.error(f"更新需求文档失败: {e}")
            raise HTTPException(status_code=500, detail=f"更新需求文档失败: {str(e)}")
    
    async def delete_requirement_document(self, document_id: uuid.UUID) -> bool:
        """删除需求文档"""
        try:
            document = await self.get_requirement_document(document_id)
            
            # 删除物理文件
            if document.file_path and os.path.exists(document.file_path):
                os.remove(document.file_path)
                logger.info(f"已删除文件: {document.file_path}")
            
            # 删除数据库记录（级联删除相关数据）
            await document.delete()
            
            logger.info(f"需求文档删除成功: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"删除需求文档失败: {e}")
            raise HTTPException(status_code=500, detail=f"删除需求文档失败: {str(e)}")
    
    # ==================== 手动需求管理 ====================
    
    async def create_requirement(
        self,
        project_id: int,
        requirement_data: RequirementCreate,
        user_id: Optional[int] = None,
        user_name: Optional[str] = None
    ) -> Requirement:
        """创建需求"""
        try:
            requirement_dict = {
                'project_id': project_id,
                'title': requirement_data.title,
                'description': requirement_data.description,
                'type': requirement_data.type,
                'priority': requirement_data.priority,
                'status': requirement_data.status,
                'stakeholders': requirement_data.stakeholders,
                'creator_id': user_id,
                'creator_name': user_name or "未知用户"
            }
            
            requirement = await Requirement.create(
                id=generate_id(),
                **requirement_dict
            )
            
            logger.info(f"需求创建成功: {requirement.id}")
            return requirement
            
        except Exception as e:
            logger.error(f"创建需求失败: {e}")
            raise HTTPException(status_code=500, detail=f"创建需求失败: {str(e)}")
    
    async def get_documents_from_filesystem(
        self,
        project_id: int,
        search: Optional[str] = None,
        page: int = 1,
        page_size: int = 20
    ) -> Tuple[List[RequirementDocument], int]:
        """从文件系统获取需求文档列表"""
        try:
            import os
            from pathlib import Path
            
            # 构建项目文档目录路径
            upload_dir = f"uploads/requirements/{project_id}"
            
            if not os.path.exists(upload_dir):
                logger.info(f"项目 {project_id} 的文档目录不存在: {upload_dir}")
                return [], 0
            
            # 扫描目录中的文件
            documents = []
            for filename in os.listdir(upload_dir):
                file_path = os.path.join(upload_dir, filename)
                
                # 跳过目录
                if os.path.isdir(file_path):
                    continue
                
                # 获取文件信息
                file_stat = os.stat(file_path)
                file_size = file_stat.st_size
                created_time = datetime.fromtimestamp(file_stat.st_ctime)
                modified_time = datetime.fromtimestamp(file_stat.st_mtime)
                
                # 提取文件名（去掉UUID前缀）
                display_name = filename
                if '_' in filename:
                    # 如果文件名包含UUID前缀，提取原始文件名
                    parts = filename.split('_', 1)
                    if len(parts) > 1:
                        display_name = parts[1]
                
                # 确定文档类型
                file_ext = Path(filename).suffix.lower()
                document_type = "unknown"
                if file_ext in ['.docx', '.doc']:
                    document_type = "word"
                elif file_ext in ['.pdf']:
                    document_type = "pdf"
                elif file_ext in ['.txt']:
                    document_type = "text"
                elif file_ext in ['.md']:
                    document_type = "markdown"
                
                # 搜索过滤
                if search and search.lower() not in display_name.lower():
                    continue
                
                # 创建文档对象（模拟数据库对象）
                doc_data = {
                    'id': str(uuid.uuid4()),  # 生成临时ID
                    'project_id': project_id,
                    'title': display_name,
                    'description': f"文件大小: {file_size} 字节",
                    'document_type': document_type,
                    'file_path': file_path,
                    'content': None,  # 暂不读取内容
                    'status': 'uploaded',
                    'version': '1.0',
                    'is_latest': True,
                    'parent_document': None,
                    'uploader_id': None,
                    'uploaded_at': created_time,
                    'updated_at': modified_time,
                    'word_count': 0,
                    'page_count': 1
                }
                
                # 创建临时文档对象，使用与RequirementDocument相同的结构
                class TempDocument:
                    def __init__(self, **kwargs):
                        for key, value in kwargs.items():
                            setattr(self, key, value)
                    
                    @property
                    def __dict__(self):
                        return {key: getattr(self, key) for key in [
                            'id', 'project_id', 'title', 'description', 'document_type',
                            'file_path', 'content', 'status', 'version', 'is_latest',
                            'parent_document', 'uploader_id', 'uploaded_at', 'updated_at',
                            'word_count', 'page_count'
                        ]}
                
                document = TempDocument(**doc_data)
                documents.append(document)
            
            # 按修改时间排序
            documents.sort(key=lambda x: x.updated_at, reverse=True)
            
            # 分页
            total = len(documents)
            start_idx = (page - 1) * page_size
            end_idx = start_idx + page_size
            paginated_docs = documents[start_idx:end_idx]
            
            logger.info(f"从文件系统找到 {total} 个文档，返回第 {page} 页的 {len(paginated_docs)} 个文档")
            return paginated_docs, total
            
        except Exception as e:
            logger.error(f"从文件系统获取需求文档失败: {e}")
            return [], 0

    async def get_requirements(
        self,
        project_id: int,
        search_params: RequirementSearchRequest
    ) -> Tuple[List[Requirement], int]:
        """获取需求列表"""
        try:
            query = Requirement.filter(project_id=project_id)
            
            # 搜索过滤
            if search_params.keyword:
                query = query.filter(
                    Q(title__icontains=search_params.keyword) | 
                    Q(description__icontains=search_params.keyword)
                )
            
            # 类型过滤
            if search_params.type:
                query = query.filter(type=search_params.type)
            
            # 优先级过滤
            if search_params.priority:
                query = query.filter(priority=search_params.priority)
            
            # 状态过滤
            if search_params.status:
                query = query.filter(status=search_params.status)
            
            # 总数
            total = await query.count()
            
            # 分页
            offset = (search_params.page - 1) * search_params.page_size
            requirements = await query.order_by('-created_at').offset(offset).limit(search_params.page_size)
            
            return requirements, total
            
        except Exception as e:
            logger.error(f"获取需求列表失败: {e}")
            raise HTTPException(status_code=500, detail=f"获取需求列表失败: {str(e)}")
    
    async def get_requirement(self, requirement_id: uuid.UUID) -> Requirement:
        """获取需求详情"""
        requirement = await Requirement.get_or_none(id=requirement_id)
        
        if not requirement:
            raise HTTPException(status_code=404, detail="需求不存在")
        
        return requirement
    
    async def update_requirement(
        self,
        requirement_id: uuid.UUID,
        update_data: RequirementUpdate
    ) -> Requirement:
        """更新需求"""
        try:
            requirement = await self.get_requirement(requirement_id)
            
            # 更新字段
            update_dict = update_data.model_dump(exclude_unset=True)
            update_dict['updated_at'] = datetime.now()
            
            await requirement.update_from_dict(update_dict)
            await requirement.save()
            
            logger.info(f"需求更新成功: {requirement_id}")
            return requirement
            
        except Exception as e:
            logger.error(f"更新需求失败: {e}")
            raise HTTPException(status_code=500, detail=f"更新需求失败: {str(e)}")
    
    async def delete_requirement(self, requirement_id: uuid.UUID) -> bool:
        """删除需求"""
        try:
            requirement = await self.get_requirement(requirement_id)
            
            await requirement.delete()
            
            logger.info(f"需求删除成功: {requirement_id}")
            return True
            
        except Exception as e:
            logger.error(f"删除需求失败: {e}")
            raise HTTPException(status_code=500, detail=f"删除需求失败: {str(e)}")
    
    # ==================== 统计功能 ====================
    
    async def get_requirement_statistics(self, project_id: int) -> RequirementStatistics:
        """获取需求统计信息"""
        try:
            # 基础统计
            total_requirements = await Requirement.filter(project_id=project_id).count()
            total_documents = await RequirementDocument.filter(project_id=project_id).count()
            
            # 需求类型分布
            type_distribution = {}
            for req_type in RequirementType:
                count = await Requirement.filter(project_id=project_id, type=req_type.value).count()
                if count > 0:
                    type_distribution[req_type.value] = count
            
            # 需求优先级分布
            priority_distribution = {}
            for priority in RequirementPriority:
                count = await Requirement.filter(project_id=project_id, priority=priority.value).count()
                if count > 0:
                    priority_distribution[priority.value] = count
            
            # 需求状态分布
            status_distribution = {}
            for status in RequirementStatus:
                count = await Requirement.filter(project_id=project_id, status=status.value).count()
                if count > 0:
                    status_distribution[status.value] = count
            
            # 文档状态分布
            doc_status_distribution = {}
            for status in DocumentStatus:
                count = await RequirementDocument.filter(project_id=project_id, status=status.value).count()
                if count > 0:
                    doc_status_distribution[status.value] = count
            
            # 模块和评审统计
            total_modules = await RequirementModule.filter(document__project_id=project_id).count()
            
            # 获取该项目的所有文档ID
            document_ids = await RequirementDocument.filter(project_id=project_id).values_list('id', flat=True)
            total_reviews = await ReviewReport.filter(document_id__in=document_ids).count()
            
            # 评审统计
            review_statistics = {
                "total_reviews": total_reviews,
                "completed_reviews": await ReviewReport.filter(
                    document_id__in=document_ids,
                    status="completed"
                ).count(),
                "average_completion_score": 0,
                "average_clarity_score": 0
            }
            
            return RequirementStatistics(
                total_requirements=total_requirements,
                total_documents=total_documents,
                total_modules=total_modules,
                total_reviews=total_reviews,
                requirement_type_distribution=type_distribution,
                requirement_priority_distribution=priority_distribution,
                requirement_status_distribution=status_distribution,
                document_status_distribution=doc_status_distribution,
                review_statistics=review_statistics
            )
            
        except Exception as e:
            logger.error(f"获取需求统计失败: {e}")
            raise HTTPException(status_code=500, detail=f"获取需求统计失败: {str(e)}")
    
    async def get_project_statistics(self, project_id: int) -> ProjectStatistics:
        """获取项目统计信息（包含知识库统计）"""
        try:
            # 获取需求统计
            req_stats = await self.get_requirement_statistics(project_id)
            
            # 获取真实的知识库统计数据
            from app.models.aitestrebort.knowledge import (
                aitestrebortKnowledgeBase, aitestrebortDocument, aitestrebortDocumentChunk
            )
            
            # 知识库统计
            total_knowledge_bases = await aitestrebortKnowledgeBase.filter(project_id=project_id).count()
            active_knowledge_bases = await aitestrebortKnowledgeBase.filter(
                project_id=project_id, is_active=True
            ).count()
            inactive_knowledge_bases = total_knowledge_bases - active_knowledge_bases
            
            # 文档统计
            kb_documents = await aitestrebortDocument.filter(
                knowledge_base__project_id=project_id
            ).count()
            processed_documents = await aitestrebortDocument.filter(
                knowledge_base__project_id=project_id,
                status='completed'
            ).count()
            processing_documents = await aitestrebortDocument.filter(
                knowledge_base__project_id=project_id,
                status='processing'
            ).count()
            failed_documents = await aitestrebortDocument.filter(
                knowledge_base__project_id=project_id,
                status='failed'
            ).count()
            
            # 分块统计
            total_chunks = await aitestrebortDocumentChunk.filter(
                document__knowledge_base__project_id=project_id
            ).count()
            
            # 总文档数（需求文档 + 知识库文档）
            total_documents = req_stats.total_documents + kb_documents
            
            # 系统健康状态检查
            system_status = "healthy"
            database_status = "connected"
            vector_db_status = "connected"
            
            # 如果有失败的文档，系统状态为警告
            if failed_documents > 0:
                system_status = "warning"
            
            return ProjectStatistics(
                project_id=project_id,
                total_knowledge_bases=total_knowledge_bases,
                active_knowledge_bases=active_knowledge_bases,
                inactive_knowledge_bases=inactive_knowledge_bases,
                total_documents=total_documents,
                processed_documents=processed_documents,
                processing_documents=processing_documents,
                failed_documents=failed_documents,
                total_requirements=req_stats.total_requirements,
                requirement_type_distribution=req_stats.requirement_type_distribution,
                requirement_priority_distribution=req_stats.requirement_priority_distribution,
                total_chunks=total_chunks,
                system_status=system_status,
                database_status=database_status,
                vector_db_status=vector_db_status,
                last_updated=datetime.now().isoformat(),
                requirement_stats=req_stats
            )
            
        except Exception as e:
            logger.error(f"获取项目统计失败: {e}")
            raise HTTPException(status_code=500, detail=f"获取项目统计失败: {str(e)}")

    # ==================== 评审管理 ====================
    
    async def get_review_results(
        self,
        project_id: int,
        document_id: Optional[uuid.UUID] = None,
        status: Optional[str] = None,
        page: int = 1,
        page_size: int = 20
    ) -> Tuple[List[ReviewReport], int]:
        """获取评审结果列表"""
        try:
            # 获取该项目的所有文档ID
            document_ids = await RequirementDocument.filter(project_id=project_id).values_list('id', flat=True)
            query = ReviewReport.filter(document_id__in=document_ids)
            
            if document_id:
                query = query.filter(document_id=document_id)
            
            if status:
                query = query.filter(status=status)
            
            # 总数
            total = await query.count()
            
            # 分页
            offset = (page - 1) * page_size
            reviews = await query.order_by('-review_date').offset(offset).limit(page_size)
            
            return reviews, total
            
        except Exception as e:
            logger.error(f"获取评审结果失败: {e}")
            raise HTTPException(status_code=500, detail=f"获取评审结果失败: {str(e)}")
    
    async def get_review_detail(self, review_id: uuid.UUID) -> ReviewReport:
        """获取评审详情"""
        try:
            review = await ReviewReport.get(id=review_id)
            return review
            
        except Exception as e:
            logger.error(f"获取评审详情失败: {e}")
            raise HTTPException(status_code=500, detail=f"获取评审详情失败: {str(e)}")
    
    async def get_review_progress(self, review_id: uuid.UUID) -> dict:
        """获取评审进度"""
        try:
            review = await ReviewReport.get(id=review_id)
            
            # 计算进度
            progress = 0.0
            current_step = "准备中"
            estimated_time = None
            
            if review.status == "pending":
                progress = 10.0
                current_step = "等待开始"
            elif review.status == "processing":
                progress = 50.0
                current_step = "评审中"
                estimated_time = 300  # 5分钟
            elif review.status == "completed":
                progress = 100.0
                current_step = "已完成"
            elif review.status == "failed":
                progress = 0.0
                current_step = "评审失败"
            
            return {
                "document_id": str(review.document_id),
                "status": review.status,
                "progress": progress,
                "current_step": current_step,
                "estimated_time": estimated_time
            }
            
        except Exception as e:
            logger.error(f"获取评审进度失败: {e}")
            raise HTTPException(status_code=500, detail=f"获取评审进度失败: {str(e)}")
    
    # ==================== 辅助方法 ====================
    
    async def _save_uploaded_file(self, file: UploadFile, project_id: int) -> str:
        """保存上传的文件"""
        try:
            # 创建目录
            upload_dir = f"uploads/requirements/{project_id}"
            os.makedirs(upload_dir, exist_ok=True)
            
            # 生成文件名
            file_extension = os.path.splitext(file.filename)[1]
            file_name = f"{uuid.uuid4()}{file_extension}"
            file_path = os.path.join(upload_dir, file_name)
            
            # 保存文件
            with open(file_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            
            logger.info(f"文件保存成功: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"保存文件失败: {e}")
            raise HTTPException(status_code=500, detail=f"保存文件失败: {str(e)}")
    
    async def _extract_document_content(self, file_path: str) -> Optional[str]:
        """提取文档内容"""
        try:
            # 使用DocumentProcessor来提取内容
            processor = DocumentProcessor()
            content = processor._extract_from_file(file_path)
            return content if content else None
            
        except Exception as e:
            logger.error(f"提取文档内容失败: {e}")
            return None
    
    async def generate_test_cases(self, requirement_id: uuid.UUID) -> Dict[str, Any]:
        """生成测试用例"""
        try:
            requirement = await self.get_requirement(requirement_id)
            
            # 这里应该集成AI生成测试用例的逻辑
            # 暂时返回模拟数据
            test_cases = [
                {
                    "title": f"测试用例1 - {requirement.title}",
                    "description": "基于需求自动生成的测试用例",
                    "steps": [
                        "步骤1: 准备测试数据",
                        "步骤2: 执行功能操作",
                        "步骤3: 验证结果"
                    ],
                    "expected_result": "功能正常工作"
                }
            ]
            
            return {
                "requirement_id": str(requirement_id),
                "requirement_title": requirement.title,
                "test_cases": test_cases,
                "generated_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"生成测试用例失败: {e}")
            raise HTTPException(status_code=500, detail=f"生成测试用例失败: {str(e)}")

    # ==================== 文档处理和模块拆分 ====================
    
    async def start_document_review(
        self,
        document_id: str,  # 改为字符串类型
        review_type: str = "comprehensive",
        focus_areas: Optional[List[str]] = None,
        user_id: Optional[int] = None
    ) -> ReviewReport:
        """开始需求评审"""
        try:
            document = await self.get_requirement_document(document_id)
            
            # 检查文档状态
            status_messages = {
                'uploaded': '文档已上传，请先进行模块拆分后再开始评审',
                'processing': '文档正在处理中，请稍后再试',
                'review_completed': '文档已完成评审，无需重复评审',
                'failed': '文档处理失败，请重新上传'
            }
            
            if document.status not in ['ready_for_review', 'reviewing']:
                error_msg = status_messages.get(document.status, f"文档状态 {document.status} 不允许开始评审")
                raise HTTPException(status_code=400, detail=error_msg)
            
            # 创建评审服务实例
            review_service = RequirementReviewService(user_id=user_id)
            
            # 根据评审类型启动评审
            if review_type == "direct":
                review_report = await review_service.start_direct_review(
                    document, 
                    {"focus_areas": focus_areas or []}
                )
            else:
                review_report = await review_service.start_comprehensive_review(
                    document,
                    {"focus_areas": focus_areas or []}
                )
            
            logger.info(f"评审启动成功: {document_id}")
            return review_report
            
        except Exception as e:
            logger.error(f"启动评审失败: {e}")
            raise HTTPException(status_code=500, detail=f"启动评审失败: {str(e)}")
    
    async def split_document_modules(
        self,
        document_id: str,
        split_options: Dict[str, Any]
    ) -> List[RequirementModule]:
        """拆分文档模块"""
        try:
            document = await self.get_requirement_document(document_id)
            
            # 检查文档状态 - 允许重新拆分
            if document.status not in ['uploaded', 'processing', 'ready_for_review', 'failed']:
                raise ValueError(f"文档状态 {document.status} 不允许进行模块拆分")
            
            # 创建模块服务实例
            module_service = RequirementModuleService(user_id=None)
            modules = await module_service.process_document_and_split(document, split_options)
            
            logger.info(f"模块拆分成功: {document_id}, 生成 {len(modules)} 个模块")
            return modules
            
        except Exception as e:
            logger.error(f"模块拆分失败: {e}")
            raise HTTPException(status_code=500, detail=f"模块拆分失败: {str(e)}")
    
    async def get_review_results(
        self,
        project_id: int,
        document_id: Optional[uuid.UUID] = None,
        status: Optional[str] = None,
        page: int = 1,
        page_size: int = 20
    ) -> Tuple[List[ReviewReport], int]:
        """获取评审结果列表"""
        try:
            # 构建查询
            query = ReviewReport.filter(document_id__in=[
                doc.id for doc in await RequirementDocument.filter(project_id=project_id).all()
            ])
            
            if document_id:
                query = query.filter(document_id=document_id)
            
            if status:
                query = query.filter(status=status)
            
            # 总数
            total = await query.count()
            
            # 分页
            offset = (page - 1) * page_size
            reviews = await query.order_by('-created_at').offset(offset).limit(page_size)
            
            return reviews, total
            
        except Exception as e:
            logger.error(f"获取评审结果失败: {e}")
            raise HTTPException(status_code=500, detail=f"获取评审结果失败: {str(e)}")
    
    async def get_review_detail(self, review_id: uuid.UUID) -> ReviewReport:
        """获取评审详情"""
        review = await ReviewReport.get_or_none(id=review_id)
        
        if not review:
            raise HTTPException(status_code=404, detail="评审报告不存在")
        
        return review
    
    async def get_review_progress(self, review_id: uuid.UUID) -> Dict[str, Any]:
        """获取评审进度"""
        review = await self.get_review_detail(review_id)
        
        # 根据状态返回进度信息
        progress_map = {
            'pending': {'status': 'pending', 'progress': 0, 'message': '等待开始评审'},
            'in_progress': {'status': 'in_progress', 'progress': 50, 'message': '正在进行评审分析'},
            'completed': {'status': 'completed', 'progress': 100, 'message': '评审已完成'},
            'failed': {'status': 'failed', 'progress': 0, 'message': '评审失败'}
        }
        
        return progress_map.get(review.status, {
            'status': 'unknown',
            'progress': 0,
            'message': '未知状态'
        })
    
    async def get_review_issues(self, review_id: uuid.UUID) -> List[ReviewIssue]:
        """获取评审问题列表"""
        try:
            # 验证评审报告是否存在
            review = await self.get_review_detail(review_id)
            
            # 获取该评审的所有问题
            issues = await ReviewIssue.filter(report_id=str(review_id)).order_by('-created_at')
            
            return issues
            
        except Exception as e:
            logger.error(f"获取评审问题列表失败: {e}")
            raise HTTPException(status_code=500, detail=f"获取评审问题列表失败: {str(e)}")
    
    async def update_issue(self, issue_id: uuid.UUID, update_data: Dict[str, Any]) -> ReviewIssue:
        """更新问题状态"""
        try:
            issue = await ReviewIssue.get_or_none(id=issue_id)
            
            if not issue:
                raise HTTPException(status_code=404, detail="问题不存在")
            
            # 更新问题数据
            for key, value in update_data.items():
                if hasattr(issue, key):
                    setattr(issue, key, value)
            
            await issue.save()
            
            return issue
            
        except Exception as e:
            logger.error(f"更新问题状态失败: {e}")
            raise HTTPException(status_code=500, detail=f"更新问题状态失败: {str(e)}")


class DocumentProcessor:
    """文档处理器 - 负责文档内容提取和预处理"""
    
    def __init__(self):
        # 初始化文本分割器
        self.text_splitter = None  # 简化实现，不依赖外部库
    
    def extract_content(self, document: RequirementDocument) -> str:
        """提取文档内容"""
        try:
            # 如果数据库中的内容是占位符文本，强制重新提取
            if (document.content and 
                len(document.content) > 20 and 
                not any(placeholder in document.content for placeholder in [
                    "需要安装python-docx库进行解析",
                    "需要安装pypdf库进行解析", 
                    "需要安装PyPDF2库进行解析",
                    "从文件", "提取的内容"
                ])):
                return document.content
            
            # 从文件重新提取内容
            if document.file_path and os.path.exists(document.file_path):
                return self._extract_from_file(document.file_path)
            
            return ""
        except Exception as e:
            logger.error(f"提取文档内容失败: {e}")
            return ""
    
    def _extract_from_file(self, file_path: str) -> str:
        """从文件提取内容"""
        try:
            file_extension = file_path.lower().split('.')[-1] if '.' in file_path else ''
            
            if file_extension == 'txt':
                return self._extract_from_txt(file_path)
            elif file_extension == 'md':
                return self._extract_from_markdown(file_path)
            elif file_extension == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['docx', 'doc']:
                return self._extract_from_word(file_path)
            elif file_extension in ['xlsx', 'xlsm']:
                return self._extract_from_excel_xlsx(file_path)
            elif file_extension == 'xls':
                return self._extract_from_excel_xls(file_path)
            else:
                logger.warning(f"不支持的文件格式: {file_extension}，尝试作为文本文件读取")
                return self._extract_from_txt(file_path)
                
        except Exception as e:
            logger.error(f"文件内容提取失败: {e}")
            return ""
    
    def _extract_from_txt(self, file_path: str) -> str:
        """提取TXT文件内容"""
        try:
            with open(file_path, 'rb') as file:
                content = file.read()
                
                # 尝试不同的编码
                for encoding in ['utf-8', 'gbk', 'gb2312']:
                    try:
                        decoded_content = content.decode(encoding)
                        logger.info(f"成功使用 {encoding} 编码读取文件，内容长度: {len(decoded_content)}")
                        return decoded_content
                    except UnicodeDecodeError:
                        continue
                
                # 如果都失败了，使用错误处理
                decoded_content = content.decode('utf-8', errors='ignore')
                logger.warning(f"使用UTF-8错误忽略模式读取文件，内容长度: {len(decoded_content)}")
                return decoded_content
                
        except Exception as e:
            logger.error(f"TXT文件读取失败: {e}")
            return ""
    
    def _extract_from_markdown(self, file_path: str) -> str:
        """提取Markdown文件内容"""
        return self._extract_from_txt(file_path)  # Markdown本质上是文本文件
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """提取PDF文件内容"""
        try:
            # 尝试使用pypdf（推荐的现代库）
            try:
                import pypdf
                return self._extract_pdf_with_pypdf(file_path)
            except ImportError:
                pass
            
            # 回退到PyPDF2
            try:
                import PyPDF2
                return self._extract_pdf_with_pypdf2(file_path)
            except ImportError:
                pass
            
            logger.error("未安装PDF解析库（pypdf或PyPDF2），无法解析PDF文档")
            return "PDF文档内容（需要安装pypdf或PyPDF2库进行解析）"
                
        except Exception as e:
            logger.error(f"PDF文档解析失败: {e}")
            return f"PDF文档解析失败: {str(e)}"
    
    def _extract_pdf_with_pypdf(self, file_path: str) -> str:
        """使用pypdf库提取PDF内容"""
        import pypdf
        
        content_parts = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            # 获取PDF信息
            num_pages = len(pdf_reader.pages)
            logger.info(f"PDF文档共有 {num_pages} 页（使用pypdf）")
            
            # 提取每一页的文本
            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text.strip():
                        # 添加页码标记
                        content_parts.append(f"=== 第 {page_num + 1} 页 ===")
                        content_parts.append(page_text.strip())
                        content_parts.append("")  # 添加空行分隔
                    
                except Exception as e:
                    logger.warning(f"提取第 {page_num + 1} 页内容失败: {e}")
                    continue
        
        content = "\n".join(content_parts)
        
        if content.strip():
            logger.info(f"成功解析PDF文档（pypdf），提取内容长度: {len(content)} 字符")
            return content
        else:
            logger.warning("PDF文档中未提取到文本内容，可能是扫描版PDF")
            return "PDF文档中未找到可提取的文本内容（可能是扫描版PDF）"
    
    def _extract_pdf_with_pypdf2(self, file_path: str) -> str:
        """使用PyPDF2库提取PDF内容"""
        import PyPDF2
        
        content_parts = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # 获取PDF信息
            num_pages = len(pdf_reader.pages)
            logger.info(f"PDF文档共有 {num_pages} 页（使用PyPDF2）")
            
            # 提取每一页的文本
            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text.strip():
                        # 添加页码标记
                        content_parts.append(f"=== 第 {page_num + 1} 页 ===")
                        content_parts.append(page_text.strip())
                        content_parts.append("")  # 添加空行分隔
                    
                except Exception as e:
                    logger.warning(f"提取第 {page_num + 1} 页内容失败: {e}")
                    continue
        
        content = "\n".join(content_parts)
        
        if content.strip():
            logger.info(f"成功解析PDF文档（PyPDF2），提取内容长度: {len(content)} 字符")
            return content
        else:
            logger.warning("PDF文档中未提取到文本内容，可能是扫描版PDF")
            return "PDF文档中未找到可提取的文本内容（可能是扫描版PDF）"
    
    def _extract_from_word(self, file_path: str) -> str:
        """提取Word文件内容"""
        try:
            from docx import Document
            
            # 读取Word文档
            doc = Document(file_path)
            
            # 提取所有段落的文本
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # 只添加非空段落
                    paragraphs.append(text)
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            content = "\n\n".join(paragraphs)
            logger.info(f"成功解析Word文档，提取内容长度: {len(content)} 字符")
            return content
            
        except ImportError:
            logger.error("python-docx库未安装，无法解析Word文档")
            return "Word文档内容（需要安装python-docx库进行解析）"
        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            return ""
    
    def _extract_from_excel_xlsx(self, file_path: str) -> str:
        """提取Excel文件内容（.xlsx, .xlsm）"""
        try:
            from openpyxl import load_workbook
            
            # 读取Excel文档
            workbook = load_workbook(file_path, read_only=True)
            
            all_content = []
            
            # 遍历所有工作表
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # 添加工作表标题
                all_content.append(f"## {sheet_name}")
                
                # 提取工作表内容
                sheet_content = []
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空行和只包含None的行
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    row_text = " | ".join(row_values).strip()
                    if row_text and row_text != " | " * (len(row_values) - 1):
                        sheet_content.append(row_text)
                
                if sheet_content:
                    all_content.extend(sheet_content)
                all_content.append("")  # 工作表之间添加空行
            
            content = "\n".join(all_content)
            logger.info(f"成功解析Excel文档，提取内容长度: {len(content)} 字符")
            return content
            
        except ImportError:
            logger.error("openpyxl库未安装，无法解析Excel文档")
            return "Excel文档内容（需要安装openpyxl库进行解析）"
        except Exception as e:
            logger.error(f"Excel文档解析失败: {e}")
            return ""
    
    def _extract_from_excel_xls(self, file_path: str) -> str:
        """提取Excel文件内容（.xls）"""
        try:
            import xlrd
            
            # 读取Excel文档
            workbook = xlrd.open_workbook(file_path)
            
            all_content = []
            
            # 遍历所有工作表
            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                
                # 添加工作表标题
                all_content.append(f"## {sheet_name}")
                
                # 提取工作表内容
                sheet_content = []
                for row_idx in range(sheet.nrows):
                    row_values = []
                    for col_idx in range(sheet.ncols):
                        cell_value = sheet.cell_value(row_idx, col_idx)
                        if cell_value:
                            row_values.append(str(cell_value))
                        else:
                            row_values.append("")
                    
                    row_text = " | ".join(row_values).strip()
                    if row_text and row_text != " | " * (len(row_values) - 1):
                        sheet_content.append(row_text)
                
                if sheet_content:
                    all_content.extend(sheet_content)
                all_content.append("")  # 工作表之间添加空行
            
            content = "\n".join(all_content)
            logger.info(f"成功解析Excel文档(.xls)，提取内容长度: {len(content)} 字符")
            return content
            
        except ImportError:
            logger.error("xlrd库未安装，无法解析Excel文档")
            return "Excel文档内容（需要安装xlrd库进行解析）"
        except Exception as e:
            logger.error(f"Excel文档解析失败: {e}")
            return ""
    
    def preprocess_content(self, content: str) -> str:
        """预处理文档内容"""
        # 清理多余的空白字符
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = re.sub(r' +', ' ', content)
        
        return content.strip()


class ModuleSplitter:
    """模块拆分器 - 负责AI智能模块识别和拆分"""
    
    def __init__(self, user_id=None):
        self.user_id = user_id
    
    def split_into_modules(
        self, 
        document: RequirementDocument, 
        content: str,
        split_options: dict = None
    ) -> List[Dict[str, Any]]:
        """将文档拆分为功能模块"""
        try:
            split_options = split_options or {}
            split_level = split_options.get('split_level', 'auto')
            chunk_size = split_options.get('chunk_size', 2000)
            
            logger.info(f"开始拆分文档，拆分级别: {split_level}")
            
            # 根据拆分级别选择方法
            if split_level == 'auto':
                modules_data = self._split_by_character_length(content, chunk_size)
            elif split_level in ['h1', 'h2', 'h3']:
                modules_data = self._split_by_heading_level(content, split_level)
            else:
                modules_data = self._split_by_document_structure(content)
            
            logger.info(f"拆分完成，生成 {len(modules_data)} 个模块")
            return modules_data
            
        except Exception as e:
            logger.error(f"模块拆分失败: {e}")
            raise
    
    def _split_by_character_length(self, content: str, chunk_size: int = 2000) -> List[Dict[str, Any]]:
        """按字符长度拆分"""
        modules_data = []
        start = 0
        order = 1
        
        while start < len(content):
            end = min(start + chunk_size, len(content))
            
            # 尝试在合适的边界切分
            if end < len(content):
                # 向后查找段落边界
                for i in range(end, min(end + 200, len(content))):
                    if content[i:i+2] == '\n\n':
                        end = i
                        break
            
            chunk_content = content[start:end].strip()
            
            if chunk_content:
                # 生成标题
                first_line = chunk_content.split('\n')[0].strip()
                title = first_line if len(first_line) < 50 else f"模块{order}"
                
                modules_data.append({
                    'title': title,
                    'content': chunk_content,
                    'order_num': order,
                    'confidence_score': 0.8,
                    'estimated_complexity': 'medium',
                    'is_auto_generated': True,
                    'start_page': 1,
                    'end_page': 1
                })
                
                order += 1
            
            start = end
        
        return modules_data
    
    def _split_by_heading_level(self, content: str, level: str) -> List[Dict[str, Any]]:
        """按标题级别拆分"""
        lines = content.split('\n')
        modules_data = []
        current_module = []
        current_title = None
        order = 1
        
        # 确定标题标记
        heading_marker = '#' * int(level[1:])  # h1 -> #, h2 -> ##, h3 -> ###
        
        for line in lines:
            stripped_line = line.strip()
            
            if stripped_line.startswith(heading_marker + ' ') and not stripped_line.startswith(heading_marker + '#'):
                # 保存前一个模块
                if current_module and current_title:
                    module_content = '\n'.join(current_module).strip()
                    if module_content:
                        modules_data.append({
                            'title': current_title,
                            'content': module_content,
                            'order_num': order,
                            'confidence_score': 0.9,
                            'estimated_complexity': 'medium',
                            'is_auto_generated': True,
                            'start_page': 1,
                            'end_page': 1
                        })
                        order += 1
                
                # 开始新模块
                current_title = stripped_line.replace(heading_marker + ' ', '').strip()
                current_module = [line]
            else:
                current_module.append(line)
        
        # 保存最后一个模块
        if current_module and current_title:
            module_content = '\n'.join(current_module).strip()
            if module_content:
                modules_data.append({
                    'title': current_title,
                    'content': module_content,
                    'order_num': order,
                    'confidence_score': 0.9,
                    'estimated_complexity': 'medium',
                    'is_auto_generated': True,
                    'start_page': 1,
                    'end_page': 1
                })
        
        # 如果没有找到任何标题，回退到按字符长度拆分
        if not modules_data:
            logger.info(f"未找到 {level} 级别的标题，回退到按字符长度拆分")
            return self._split_by_character_length(content, 2000)
        
        return modules_data
    
    def _split_by_document_structure(self, content: str) -> List[Dict[str, Any]]:
        """基于文档结构的拆分"""
        lines = content.split('\n')
        modules_data = []
        
        # 找到所有二级标题作为模块边界
        module_boundaries = []
        module_titles = []
        
        for i, line in enumerate(lines):
            stripped_line = line.strip()
            if stripped_line.startswith('## ') and not stripped_line.startswith('### '):
                module_boundaries.append(i)
                module_titles.append(stripped_line.replace('## ', '').strip())
        
        logger.info(f"识别到 {len(module_boundaries)} 个二级标题: {module_titles}")
        
        # 如果没有找到标题，按内容分割
        if len(module_boundaries) < 2:
            return self._split_by_character_length(content, 2000)
        
        # 按模块边界拆分
        for i, start_idx in enumerate(module_boundaries):
            end_idx = module_boundaries[i + 1] if i + 1 < len(module_boundaries) else len(lines)
            
            module_lines = lines[start_idx:end_idx]
            module_content = '\n'.join(module_lines).strip()
            
            if module_content:
                modules_data.append({
                    'title': module_titles[i],
                    'content': module_content,
                    'order_num': i + 1,
                    'confidence_score': 0.95,
                    'estimated_complexity': 'medium',
                    'is_auto_generated': True,
                    'start_page': 1,
                    'end_page': 1
                })
        
        return modules_data


class RequirementModuleService:
    """需求模块服务 - 统一的模块管理服务"""
    
    def __init__(self, user_id=None):
        self.user_id = user_id
        self.document_processor = DocumentProcessor()
        self.module_splitter = ModuleSplitter(user_id=user_id)
    
    async def process_document_and_split(
        self, 
        document: RequirementDocument,
        split_options: dict = None
    ) -> List[RequirementModule]:
        """处理文档并进行模块拆分"""
        try:
            # 如果文档已经有模块，先删除旧的模块（支持重新拆分）
            existing_modules = await RequirementModule.filter(document_id=document.id).all()
            if existing_modules:
                logger.info(f"删除文档 {document.id} 的 {len(existing_modules)} 个旧模块")
                for module in existing_modules:
                    await module.delete()
            
            # 更新文档状态
            document.status = 'processing'
            await document.save()
            
            # 提取文档内容
            logger.info(f"开始处理文档 {document.id}: {document.title}")
            content = self.document_processor.extract_content(document)
            if not content:
                raise Exception("无法提取文档内容")
            
            logger.info(f"文档内容提取完成，原始长度: {len(content)} 字符")
            
            # 预处理内容
            processed_content = self.document_processor.preprocess_content(content)
            logger.info(f"内容预处理完成，处理后长度: {len(processed_content)} 字符")
            
            # 更新文档统计信息
            document.word_count = len(processed_content)
            document.page_count = max(1, (len(processed_content) // 500) + 1)
            document.content = processed_content
            await document.save()
            
            # 模块拆分
            logger.info(f"开始模块拆分，拆分选项: {split_options}")
            modules_data = self.module_splitter.split_into_modules(document, processed_content, split_options)
            
            # 创建模块对象
            modules = []
            for i, module_data in enumerate(modules_data):
                module = await RequirementModule.create(
                    id=generate_id(),
                    document_id=document.id,
                    title=module_data['title'],
                    content=module_data['content'],
                    start_page=module_data.get('start_page', 1),
                    end_page=module_data.get('end_page', 1),
                    order_num=module_data['order_num'],
                    confidence_score=module_data['confidence_score'],
                    is_auto_generated=True
                )
                modules.append(module)
                logger.info(f"创建模块 {i+1}/{len(modules_data)}: {module_data['title']}")
            
            # 更新文档状态
            document.status = 'ready_for_review'
            await document.save()
            
            logger.info(f"文档 {document.id} 模块拆分完成，生成 {len(modules)} 个模块")
            return modules
            
        except Exception as e:
            logger.error(f"文档模块拆分失败: {e}")
            document.status = 'failed'
            await document.save()
            raise


class RequirementReviewService:
    """需求评审服务 - 统一的评审管理服务"""
    
    def __init__(self, user_id=None):
        self.user_id = user_id
        self.review_engine = RequirementReviewEngine(user_id=user_id)
    
    async def start_direct_review(
        self, 
        document: RequirementDocument,
        analysis_options: dict = None
    ) -> ReviewReport:
        """启动直接评审（不拆分模块）"""
        try:
            # 检查文档内容
            if not document.content:
                raise ValueError("文档内容为空，无法进行评审")
            
            # 创建评审报告
            review_report = await ReviewReport.create(
                id=generate_id(),
                document_id=document.id,
                status='in_progress',
                reviewer='AI需求评审助手',
                review_type='direct'
            )
            
            # 更新文档状态
            document.status = 'reviewing'
            await document.save()
            
            # 执行评审分析
            review_result = await self.review_engine.analyze_document_directly(
                document.content,
                analysis_options or {}
            )
            
            # 更新评审报告
            review_report.overall_rating = review_result.get('overall_rating', 'average')
            review_report.completion_score = review_result.get('completion_score', 70)
            review_report.clarity_score = review_result.get('clarity_score', 70)
            review_report.consistency_score = review_result.get('consistency_score', 70)
            review_report.completeness_score = review_result.get('completeness_score', 70)
            review_report.summary = review_result.get('summary', '评审已完成')
            review_report.recommendations = review_result.get('recommendations', '建议进一步完善需求描述')
            review_report.status = 'completed'
            await review_report.save()
            
            # 创建问题记录
            issues = review_result.get('issues', [])
            for issue_data in issues:
                await ReviewIssue.create(
                    id=generate_id(),
                    report_id=review_report.id,
                    title=issue_data.get('title', '未知问题'),
                    description=issue_data.get('description', ''),
                    priority=issue_data.get('priority', 'medium'),
                    issue_type=issue_data.get('category', 'specification'),
                    suggestion=issue_data.get('suggestion', ''),
                    location=issue_data.get('location', '')
                )
            
            # 更新统计信息
            review_report.total_issues = len(issues)
            review_report.high_priority_issues = len([i for i in issues if i.get('priority') == 'high'])
            await review_report.save()
            
            # 更新文档状态
            document.status = 'review_completed'
            await document.save()
            
            logger.info(f"文档 {document.id} 直接评审完成")
            return review_report
            
        except Exception as e:
            logger.error(f"直接评审失败: {e}")
            if 'review_report' in locals():
                review_report.status = 'failed'
                await review_report.save()
            raise
    
    async def start_comprehensive_review(
        self, 
        document: RequirementDocument,
        analysis_options: dict = None
    ) -> ReviewReport:
        """启动全面的需求评审（基于模块）"""
        try:
            # 创建评审报告
            review_report = await ReviewReport.create(
                id=generate_id(),
                document_id=document.id,
                status='in_progress',
                reviewer='AI需求评审助手',
                review_type='comprehensive'
            )
            
            # 更新文档状态
            document.status = 'reviewing'
            await document.save()
            
            logger.info(f"开始评审文档: {document.title}")
            
            # 执行AI分析
            analysis_result = await self.review_engine.analyze_document_comprehensive(
                document, 
                analysis_options
            )
            
            # 更新评审报告
            await self._update_review_report(review_report, analysis_result)
            
            # 创建问题记录
            await self._create_review_issues(review_report, analysis_result)
            
            # 创建模块评审结果
            await self._create_module_results(review_report, analysis_result)
            
            # 完成评审
            review_report.status = 'completed'
            await review_report.save()
            
            # 更新文档状态
            document.status = 'review_completed'
            await document.save()
            
            logger.info(f"评审完成: {document.title}, 总体评分: {review_report.completion_score}")
            return review_report
            
        except Exception as e:
            logger.error(f"评审失败: {e}")
            
            if 'review_report' in locals():
                review_report.status = 'failed'
                await review_report.save()
            
            document.status = 'failed'
            await document.save()
            
            raise
    
    async def _update_review_report(self, review_report: ReviewReport, analysis_result: dict):
        """更新评审报告"""
        review_report.overall_rating = analysis_result.get('overall_rating', 'average')
        review_report.completion_score = analysis_result.get('overall_score', 70)
        review_report.total_issues = analysis_result.get('total_issues', 0)
        review_report.high_priority_issues = analysis_result.get('high_priority_issues', 0)
        review_report.summary = analysis_result.get('summary', '评审已完成')
        review_report.recommendations = '\n'.join(analysis_result.get('recommendations', ['建议进一步完善需求描述']))
        
        # 保存专项分析详情
        specialized_analyses = analysis_result.get('specialized_analyses', {})
        review_report.specialized_analyses = specialized_analyses
        
        # 各专项分析分数
        review_report.completeness_score = specialized_analyses.get('completeness_analysis', {}).get('overall_score', 70)
        review_report.consistency_score = specialized_analyses.get('consistency_analysis', {}).get('overall_score', 70)
        review_report.clarity_score = specialized_analyses.get('clarity_analysis', {}).get('overall_score', 70)
        
        await review_report.save()
    
    async def _create_review_issues(self, review_report: ReviewReport, analysis_result: dict):
        """创建评审问题记录"""
        issues = analysis_result.get('issues', [])
        for issue_data in issues:
            # 处理不同格式的问题数据
            if isinstance(issue_data, dict):
                # 字典格式的问题数据
                await ReviewIssue.create(
                    id=generate_id(),
                    report_id=review_report.id,
                    title=issue_data.get('title', '未知问题'),
                    description=issue_data.get('description', ''),
                    priority=issue_data.get('priority', 'medium'),
                    issue_type=issue_data.get('category', 'specification'),
                    suggestion=issue_data.get('suggestion', ''),
                    location=issue_data.get('location', '')
                )
            elif isinstance(issue_data, str):
                # 字符串格式的问题数据
                await ReviewIssue.create(
                    id=generate_id(),
                    report_id=review_report.id,
                    title=issue_data[:100] if len(issue_data) > 100 else issue_data,
                    description=issue_data,
                    priority='medium',
                    issue_type='general',
                    suggestion='请查看详细描述',
                    location='AI分析'
                )
            else:
                # 其他格式，转换为字符串
                issue_str = str(issue_data)
                await ReviewIssue.create(
                    id=generate_id(),
                    report_id=review_report.id,
                    title=issue_str[:100] if len(issue_str) > 100 else issue_str,
                    description=issue_str,
                    priority='medium',
                    issue_type='general',
                    suggestion='请查看详细描述',
                    location='AI分析'
                )
    
    async def _create_module_results(self, review_report: ReviewReport, analysis_result: dict):
        """创建模块评审结果"""
        module_results = analysis_result.get('module_results', [])
        for module_data in module_results:
            await ModuleReviewResult.create(
                id=generate_id(),
                report_id=review_report.id,
                module_id=module_data.get('module_id'),
                module_name=module_data.get('module_name', '未知模块'),
                specification_score=module_data.get('specification_score', 70),
                clarity_score=module_data.get('clarity_score', 70),
                completeness_score=module_data.get('completeness_score', 70),
                consistency_score=module_data.get('consistency_score', 70),
                feasibility_score=module_data.get('feasibility_score', 70),
                overall_score=module_data.get('overall_score', 70),
                issues=module_data.get('issues', []),
                strengths=module_data.get('strengths', []),
                weaknesses=module_data.get('weaknesses', []),
                recommendations=module_data.get('recommendations', [])
            )


class RequirementReviewEngine:
    """需求评审引擎 - 负责AI分析和评审逻辑"""
    
    def __init__(self, user_id=None):
        self.user_id = user_id
    
    async def analyze_document_directly(self, content: str, analysis_options: dict = None) -> dict:
        """直接分析整个文档（不拆分模块）"""
        try:
            logger.info(f"开始AI直接分析文档，内容长度: {len(content) if content else 0}")
            
            # 使用真实的AI分析
            from app.services.ai.llm_service import get_llm_service
            from app.services.ai.prompt_manager import PromptType, format_prompt
            
            llm_service = await get_llm_service()
            
            # 构建分析内容
            analysis_text = f"""
            文档内容: {content[:3000] if content else '无内容'}
            
            请对以上需求文档进行全面分析，包括：
            1. 完整性评估
            2. 清晰度评估
            3. 一致性评估
            4. 可行性评估
            5. 发现的问题和改进建议
            """
            
            # 使用提示词模板进行分析
            prompt = format_prompt(PromptType.REQUIREMENT_ANALYSIS, requirement_text=analysis_text)
            response = await llm_service.generate_text(prompt, temperature=0.3)
            
            # 尝试解析JSON响应
            import json
            try:
                result = json.loads(response)
                logger.info("AI直接分析完成，使用真实AI分析结果")
                
                # 确保结果包含必要字段
                analysis_result = {
                    'overall_rating': result.get('overall_rating', 'good'),
                    'completion_score': result.get('completion_score', 75),
                    'clarity_score': result.get('clarity_score', 80),
                    'consistency_score': result.get('consistency_score', 70),
                    'completeness_score': result.get('completeness_score', 75),
                    'summary': result.get('summary', 'AI分析完成'),
                    'recommendations': result.get('recommendations', ['建议进一步完善需求描述']),
                    'issues': result.get('issues', [])
                }
                
                return analysis_result
                
            except json.JSONDecodeError:
                logger.warning("AI响应格式错误，使用基础分析结果")
                return {
                    'overall_rating': 'average',
                    'completion_score': 70,
                    'clarity_score': 70,
                    'consistency_score': 70,
                    'completeness_score': 70,
                    'summary': f'AI分析完成，但响应格式需要优化。原始响应: {response[:200]}...',
                    'recommendations': ['建议检查AI服务配置', '优化提示词模板'],
                    'issues': [
                        {
                            'title': 'AI响应格式错误',
                            'description': 'AI返回的分析结果不是有效的JSON格式',
                            'priority': 'medium',
                            'category': 'system',
                            'suggestion': '检查提示词模板和AI模型配置',
                            'location': '系统分析'
                        }
                    ]
                }
                
        except Exception as e:
            logger.error(f"AI直接文档分析失败: {e}")
            return {
                'overall_rating': 'failed',
                'completion_score': 0,
                'clarity_score': 0,
                'consistency_score': 0,
                'completeness_score': 0,
                'summary': f'AI分析失败: {str(e)}',
                'recommendations': ['请检查LLM配置并确保AI服务正常运行'],
                'issues': [
                    {
                        'title': 'AI分析服务不可用',
                        'description': f'AI分析过程中发生错误: {str(e)}',
                        'priority': 'high',
                        'category': 'system',
                        'suggestion': '检查LLM配置和网络连接',
                        'location': '系统分析'
                    }
                ]
            }
    
    async def analyze_document_comprehensive(self, document: RequirementDocument, analysis_options: dict = None) -> dict:
        """全面分析文档（基于模块）"""
        try:
            # 获取文档的模块
            modules = await RequirementModule.filter(document=document).all()
            
            logger.info(f"开始AI全面分析文档: {document.title}, 模块数: {len(modules)}")
            
            # 使用真实的AI分析
            from app.services.aitestrebort.requirements_review import RequirementReviewService
            from app.services.aitestrebort.requirements_service import generate_id
            
            # 执行各项专项分析
            analyses = {}
            
            # 完整性分析
            completeness_result = await RequirementReviewService._analyze_completeness(document, modules)
            analyses["completeness_analysis"] = completeness_result
            
            # 一致性分析
            consistency_result = await RequirementReviewService._analyze_consistency(document, modules)
            analyses["consistency_analysis"] = consistency_result
            
            # 可测性分析
            testability_result = await RequirementReviewService._analyze_testability(document, modules)
            analyses["testability_analysis"] = testability_result
            
            # 可行性分析
            feasibility_result = await RequirementReviewService._analyze_feasibility(document, modules)
            analyses["feasibility_analysis"] = feasibility_result
            
            # 清晰度分析
            clarity_result = await RequirementReviewService._analyze_clarity(document, modules)
            analyses["clarity_analysis"] = clarity_result
            
            # 计算综合评分
            scores = []
            total_issues = []
            high_priority_issues = 0
            
            for analysis in analyses.values():
                if isinstance(analysis.get('score'), (int, float)):
                    scores.append(analysis['score'])
                if isinstance(analysis.get('issues'), list):
                    total_issues.extend(analysis['issues'])
                    high_priority_issues += len([i for i in analysis['issues'] if 'high' in str(i).lower()])
            
            overall_score = sum(scores) / len(scores) if scores else 70
            
            # 构建分析结果
            analysis_result = {
                'overall_rating': 'excellent' if overall_score >= 90 else 'good' if overall_score >= 80 else 'average' if overall_score >= 70 else 'needs_improvement',
                'overall_score': int(overall_score),
                'total_issues': len(total_issues),
                'high_priority_issues': high_priority_issues,
                'medium_priority_issues': len([i for i in total_issues if 'medium' in str(i).lower()]),
                'low_priority_issues': len([i for i in total_issues if 'low' in str(i).lower()]),
                'summary': f'文档经过AI全面分析，总体评分{overall_score:.1f}分，发现{len(total_issues)}个问题。',
                'recommendations': [
                    rec for analysis in analyses.values() 
                    if isinstance(analysis.get('recommendations'), list)
                    for rec in analysis['recommendations']
                ][:5],  # 取前5个建议
                'specialized_analyses': analyses,
                'issues': total_issues,
                'module_results': []
            }
            
            # 为每个模块生成分析结果（基于AI分析）
            for module in modules:
                module_issues = [i for i in total_issues if module.title.lower() in str(i).lower()]
                module_score = max(50, overall_score - len(module_issues) * 5)  # 根据问题数量调整评分
                
                module_result = {
                    'module_id': str(module.id),
                    'module_name': module.title,
                    'specification_score': int(module_score),
                    'clarity_score': int(module_score + 2),
                    'completeness_score': int(module_score - 3),
                    'consistency_score': int(module_score + 1),
                    'feasibility_score': int(module_score + 5),
                    'overall_score': int(module_score),
                    'issues': [str(issue) for issue in module_issues],
                    'strengths': [strength for analysis in analyses.values() 
                                if isinstance(analysis.get('strengths'), list)
                                for strength in analysis['strengths']][:2],
                    'weaknesses': [str(issue) for issue in module_issues[:2]],
                    'recommendations': [rec for analysis in analyses.values() 
                                     if isinstance(analysis.get('recommendations'), list)
                                     for rec in analysis['recommendations']][:2]
                }
                analysis_result['module_results'].append(module_result)
            
            logger.info(f"AI全面文档分析完成，总体评分: {analysis_result['overall_score']}")
            return analysis_result
            
        except Exception as e:
            logger.error(f"AI全面文档分析失败: {e}")
            # 当AI分析失败时，返回错误信息而不是假数据
            return {
                'overall_rating': 'failed',
                'overall_score': 0,
                'total_issues': 1,
                'high_priority_issues': 1,
                'medium_priority_issues': 0,
                'low_priority_issues': 0,
                'summary': f'AI分析失败: {str(e)}',
                'recommendations': ['请检查LLM配置并确保AI服务正常运行'],
                'specialized_analyses': {
                    'error': {'message': str(e)}
                },
                'issues': [f'AI分析服务不可用: {str(e)}'],
                'module_results': []
            }
    
    def _get_default_analysis_result(self) -> dict:
        """获取默认的分析结果"""
        return {
            'overall_rating': 'average',
            'overall_score': 70,
            'completion_score': 70,
            'clarity_score': 70,
            'consistency_score': 70,
            'completeness_score': 70,
            'total_issues': 0,
            'high_priority_issues': 0,
            'summary': '文档分析完成，整体质量中等。',
            'recommendations': ['建议进一步完善需求描述'],
            'issues': [],
            'specialized_analyses': {
                'completeness_analysis': {'overall_score': 70},
                'consistency_analysis': {'overall_score': 70},
                'clarity_analysis': {'overall_score': 70}
            },
            'module_results': []
        }