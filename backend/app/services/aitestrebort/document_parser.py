"""
统一文档解析服务
支持多种文档格式的解析和结构化处理
"""
import os
import logging
from typing import Dict, Any, Optional, List
from abc import ABC, abstractmethod
from datetime import datetime
import hashlib

logger = logging.getLogger(__name__)


class BaseParser(ABC):
    """文档解析器基类"""
    
    @abstractmethod
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析文档
        Returns:
            {
                'content': str,  # 文档内容
                'metadata': dict,  # 元数据
                'structure': dict  # 结构信息
            }
        """
        pass
    
    def _extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取基础元数据"""
        stat = os.stat(file_path)
        return {
            'file_name': os.path.basename(file_path),
            'file_size': stat.st_size,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'file_hash': self._calculate_hash(file_path)
        }
    
    def _calculate_hash(self, file_path: str) -> str:
        """计算文件哈希"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()


class PDFParser(BaseParser):
    """PDF文档解析器"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            import pdfplumber
            
            content_parts = []
            metadata = self._extract_metadata(file_path)
            tables = []
            
            with pdfplumber.open(file_path) as pdf:
                metadata['page_count'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取文本
                    text = page.extract_text()
                    if text:
                        content_parts.append(f"[Page {page_num}]\n{text}")
                    
                    # 提取表格
                    page_tables = page.extract_tables()
                    if page_tables:
                        for table in page_tables:
                            tables.append({
                                'page': page_num,
                                'data': table
                            })
            
            content = "\n\n".join(content_parts)
            metadata['word_count'] = len(content.split())
            
            return {
                'content': content,
                'metadata': metadata,
                'structure': {
                    'tables': tables,
                    'page_count': metadata['page_count']
                }
            }
            
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            raise ValueError(f"PDF解析失败: {str(e)}")


class WordParser(BaseParser):
    """Word文档解析器"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            from docx import Document
            
            doc = Document(file_path)
            metadata = self._extract_metadata(file_path)
            
            # 提取文本
            content_parts = []
            tables = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            
            # 提取表格
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    'index': table_idx,
                    'data': table_data
                })
            
            content = "\n\n".join(content_parts)
            metadata['word_count'] = len(content.split())
            metadata['paragraph_count'] = len(doc.paragraphs)
            
            # 提取文档属性
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            
            return {
                'content': content,
                'metadata': metadata,
                'structure': {
                    'tables': tables,
                    'paragraph_count': metadata['paragraph_count']
                }
            }
            
        except Exception as e:
            logger.error(f"Word解析失败: {e}")
            raise ValueError(f"Word解析失败: {str(e)}")


class ExcelParser(BaseParser):
    """Excel文档解析器"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            import pandas as pd
            
            metadata = self._extract_metadata(file_path)
            
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            sheets_data = []
            content_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # 转换为文本
                sheet_text = f"[Sheet: {sheet_name}]\n"
                sheet_text += df.to_string(index=False)
                content_parts.append(sheet_text)
                
                # 保存结构化数据
                sheets_data.append({
                    'name': sheet_name,
                    'rows': len(df),
                    'columns': list(df.columns),
                    'data': df.to_dict('records')
                })
            
            content = "\n\n".join(content_parts)
            metadata['sheet_count'] = len(excel_file.sheet_names)
            
            return {
                'content': content,
                'metadata': metadata,
                'structure': {
                    'sheets': sheets_data
                }
            }
            
        except Exception as e:
            logger.error(f"Excel解析失败: {e}")
            raise ValueError(f"Excel解析失败: {str(e)}")


class MarkdownParser(BaseParser):
    """Markdown文档解析器"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            import markdown
            from bs4 import BeautifulSoup
            
            metadata = self._extract_metadata(file_path)
            
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 转换为HTML以提取结构
            html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
            soup = BeautifulSoup(html, 'html.parser')
            
            # 提取标题结构
            headings = []
            for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                headings.append({
                    'level': int(heading.name[1]),
                    'text': heading.get_text()
                })
            
            # 提取代码块
            code_blocks = []
            for code in soup.find_all('code'):
                code_blocks.append(code.get_text())
            
            metadata['word_count'] = len(md_content.split())
            metadata['heading_count'] = len(headings)
            
            return {
                'content': md_content,
                'metadata': metadata,
                'structure': {
                    'headings': headings,
                    'code_blocks': code_blocks
                }
            }
            
        except Exception as e:
            logger.error(f"Markdown解析失败: {e}")
            raise ValueError(f"Markdown解析失败: {str(e)}")


class HTMLParser(BaseParser):
    """HTML文档解析器"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            from bs4 import BeautifulSoup
            
            metadata = self._extract_metadata(file_path)
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除脚本和样式
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 提取文本
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)
            
            # 提取标题
            title = soup.find('title')
            if title:
                metadata['title'] = title.get_text()
            
            # 提取元数据
            meta_tags = soup.find_all('meta')
            for meta in meta_tags:
                if meta.get('name') == 'description':
                    metadata['description'] = meta.get('content')
                elif meta.get('name') == 'author':
                    metadata['author'] = meta.get('content')
            
            metadata['word_count'] = len(content.split())
            
            return {
                'content': content,
                'metadata': metadata,
                'structure': {
                    'has_title': bool(title),
                    'meta_count': len(meta_tags)
                }
            }
            
        except Exception as e:
            logger.error(f"HTML解析失败: {e}")
            raise ValueError(f"HTML解析失败: {str(e)}")


class TextParser(BaseParser):
    """纯文本解析器"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        try:
            metadata = self._extract_metadata(file_path)
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata['word_count'] = len(content.split())
            metadata['line_count'] = len(content.splitlines())
            
            return {
                'content': content,
                'metadata': metadata,
                'structure': {
                    'line_count': metadata['line_count']
                }
            }
            
        except Exception as e:
            logger.error(f"文本解析失败: {e}")
            raise ValueError(f"文本解析失败: {str(e)}")


class DocumentParser:
    """统一文档解析器"""
    
    # 支持的文件类型映射
    PARSERS = {
        'pdf': PDFParser,
        'docx': WordParser,
        'doc': WordParser,
        'xlsx': ExcelParser,
        'xls': ExcelParser,
        'md': MarkdownParser,
        'markdown': MarkdownParser,
        'html': HTMLParser,
        'htm': HTMLParser,
        'txt': TextParser,
        'text': TextParser
    }
    
    @classmethod
    async def parse(cls, file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        解析文档
        Args:
            file_path: 文件路径
            file_type: 文件类型（可选，如果不提供则从文件扩展名推断）
        Returns:
            {
                'content': str,
                'metadata': dict,
                'structure': dict
            }
        """
        # 推断文件类型
        if not file_type:
            _, ext = os.path.splitext(file_path)
            file_type = ext.lstrip('.').lower()
        
        # 获取对应的解析器
        parser_class = cls.PARSERS.get(file_type)
        if not parser_class:
            raise ValueError(f"不支持的文件类型: {file_type}")
        
        logger.info(f"使用 {parser_class.__name__} 解析文件: {file_path}")
        
        # 解析文档
        parser = parser_class()
        result = await parser.parse(file_path)
        
        # 添加解析器信息
        result['metadata']['parser'] = parser_class.__name__
        result['metadata']['file_type'] = file_type
        
        return result
    
    @classmethod
    def get_supported_types(cls) -> List[str]:
        """获取支持的文件类型列表"""
        return list(cls.PARSERS.keys())
    
    @classmethod
    def is_supported(cls, file_type: str) -> bool:
        """检查是否支持指定的文件类型"""
        return file_type.lower() in cls.PARSERS


# 导出
__all__ = ['DocumentParser', 'BaseParser']
